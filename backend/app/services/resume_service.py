from typing import Optional, BinaryIO
import uuid
import json
import structlog
import io
import re
import tempfile
from pathlib import Path

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func, desc
from sqlalchemy.orm import selectinload

from app.core.exceptions import NotFoundException, ValidationException
from app.core.notifier import create_notification
from app.core.s3 import upload_file, get_presigned_url, delete_file
from app.models.user import User
from app.models.job import Job
from app.models.resume import Resume, ResumeVersion, ResumeSkillGraph, GeneratedResume, ParsingStatus, FileType
from app.schemas.resume import (
    ResumeUploadResponse, ResumeResponse, ResumeAnalysisResponse,
    SkillItem, ResumeTailorResponse, ParsedResumeData,
)
from app.core.ai.gateway import get_gateway
from app.core.config import settings

logger = structlog.get_logger()

TEXT_MIN_THRESHOLD = 100


class ResumeService:
    def __init__(self, db: AsyncSession):
        self.db = db

    async def upload_resume(self, user: User, file) -> ResumeUploadResponse:
        content = await file.read()
        file_size = len(content)
        file_ext = Path(file.filename or "resume.pdf").suffix.lower()
        if file_ext == ".pdf":
            file_type = "pdf"
        elif file_ext in (".docx", ".doc"):
            file_type = "docx"
        else:
            raise ValidationException("Only PDF and DOCX files are accepted")

        if file_size > 10 * 1024 * 1024:
            raise ValidationException("File size exceeds 10MB limit")

        file_key = f"resumes/{user.id}/{uuid.uuid4()}_{file.filename}"

        resume = Resume(
            user_id=user.id,
            file_key=file_key,
            file_name=file.filename,
            file_size=file_size,
            file_type=FileType(file_type),
            parsing_status=ParsingStatus.PROCESSING,
        )
        self.db.add(resume)
        await self.db.flush()
        await self.db.refresh(resume)

        try:
            buf = io.BytesIO(content)
            await upload_file(settings.S3_BUCKET, file_key, buf, file.content_type or "application/octet-stream")
            logger.info("resume_stored_in_minio", file_key=file_key, file_size=file_size)
        except Exception as e:
            logger.error("minio_upload_failed", error=str(e))
            resume.parsing_status = ParsingStatus.FAILED
            resume.parsing_error = f"Storage failed: {str(e)}"
            await self.db.flush()
            return ResumeUploadResponse(
                id=str(resume.id),
                file_name=resume.file_name,
                file_size=resume.file_size,
                file_type=resume.file_type.value,
                parsing_status=resume.parsing_status.value,
                created_at=resume.created_at,
            )

        parsed_data, raw_text, confidence, error = await self._parse_resume_content(content, file_type)
        embedding = await self._generate_embedding(raw_text or "")

        resume.parsing_status = ParsingStatus.COMPLETED if not error else ParsingStatus.FAILED
        resume.raw_text = raw_text
        resume.parsed_data = parsed_data
        resume.parsing_confidence = confidence
        resume.parsing_error = error
        if not error:
            user.active_resume_id = resume.id
        await self.db.flush()

        if embedding:
            existing = await self.db.execute(
                select(ResumeSkillGraph).where(ResumeSkillGraph.resume_id == resume.id)
            )
            skill_graph = existing.scalar_one_or_none()
            if skill_graph:
                skill_graph.embedding = embedding
            else:
                new_graph = ResumeSkillGraph(
                    resume_id=resume.id,
                    skills=parsed_data.get("skills", []) if parsed_data else [],
                    embedding=embedding,
                )
                self.db.add(new_graph)
            await self.db.flush()

        if not error:
            try:
                from app.services.matching_service import MatchingService
                matching = MatchingService(self.db)
                # Keep optional insight generation isolated so a failure cannot
                # roll back an otherwise successful resume upload.
                async with self.db.begin_nested():
                    await matching.generate_career_insights(user.id, resume.id)
                logger.info("career_insights_generated_after_upload", resume_id=str(resume.id))
            except Exception as e:
                logger.warning("resume_post_processing_failed", error=str(e))

            # Job matching is intentionally handled by the periodic worker and
            # on-demand search. Scoring the entire catalog here makes the upload
            # request time out as the catalog grows.

        await create_notification(
            self.db, user.id, "resume_parsed",
            f"Resume {'parsed' if not error else 'parsing failed'}: {resume.file_name}",
            body=f"Extracted content with {int(confidence * 100)}% confidence"
            if not error else f"Error: {error}",
        )

        return ResumeUploadResponse(
            id=str(resume.id),
            file_name=resume.file_name,
            file_size=resume.file_size,
            file_type=resume.file_type.value,
            parsing_status=resume.parsing_status.value,
            created_at=resume.created_at,
        )

    async def _parse_resume_content(self, content: bytes, file_type: str) -> tuple[Optional[dict], Optional[str], float, Optional[str]]:
        raw_text = None
        confidence = 0.0

        try:
            if file_type == "pdf":
                raw_text = await self._extract_pdf_text(content)
            elif file_type == "docx":
                raw_text = await self._extract_docx_text(content)
        except Exception as e:
            logger.error("text_extraction_failed", error=str(e))
            return None, None, 0.0, f"Text extraction failed: {str(e)}"

        if not raw_text or len(raw_text.strip()) < TEXT_MIN_THRESHOLD:
            logger.info("low_text_quality_falling_back_to_ocr", length=len(raw_text or ""))
            try:
                if file_type == "pdf":
                    ocr_text = await self._ocr_pdf(content)
                    if ocr_text and len(ocr_text.strip()) > len((raw_text or "").strip()):
                        raw_text = ocr_text
            except Exception as e:
                logger.warning("ocr_fallback_failed", error=str(e))

        if not raw_text or len(raw_text.strip()) < TEXT_MIN_THRESHOLD:
            return None, raw_text, 0.1, "Could not extract meaningful text from the file. The file may be an image-based PDF or corrupted."

        raw_text_clean = raw_text.strip()
        text_len = len(raw_text_clean)
        confidence = min(0.95, 0.3 + (text_len / 5000) * 0.6)
        if text_len < 200:
            confidence = max(0.15, text_len / 200 * 0.5)

        try:
            parsed = await self._ai_structure_resume(raw_text_clean)
            if not parsed:
                parsed = self._fallback_parse(raw_text_clean)
            return parsed, raw_text_clean, confidence, None
        except Exception as e:
            logger.error("ai_structuring_failed", error=str(e))
            parsed = self._fallback_parse(raw_text_clean)
            return parsed, raw_text_clean, confidence * 0.7, None

    async def _extract_pdf_text(self, content: bytes) -> Optional[str]:
        import fitz
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            if text.strip():
                logger.info("pdf_extract_pymupdf_success", length=len(text))
                return text
        except Exception as e:
            logger.warning("pymupdf_failed", error=str(e))

        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                logger.info("pdf_extract_pdfplumber_success", length=len(text))
                return text
        except ImportError:
            logger.warning("pdfplumber_not_available")
        except Exception as e:
            logger.warning("pdfplumber_failed", error=str(e))

        return None

    async def _extract_docx_text(self, content: bytes) -> Optional[str]:
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(content))
        text_parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text_parts.append(" | ".join(cells))
        text = "\n".join(text_parts)
        logger.info("docx_extract_success", length=len(text))
        return text

    async def _ocr_pdf(self, content: bytes) -> Optional[str]:
        try:
            from pdf2image import convert_from_bytes
            try:
                import pytesseract
            except ImportError:
                logger.warning("pytesseract_not_available")
                return None

            images = convert_from_bytes(content, dpi=300)
            text = ""
            self._log_ocr_progress(0, len(images))
            for i, img in enumerate(images):
                page_text = pytesseract.image_to_string(img, lang="eng")
                text += page_text + "\n"
                self._log_ocr_progress(i + 1, len(images))
            logger.info("ocr_extract_success", pages=len(images), length=len(text))
            return text
        except ImportError as e:
            logger.warning("ocr_dependency_missing", error=str(e))
            return None
        except Exception as e:
            logger.warning("ocr_extract_failed", error=str(e))
            return None

    def _log_ocr_progress(self, current: int, total: int):
        if total <= 1:
            return
        if current % max(1, total // 5) == 0 or current == total:
            logger.info("ocr_progress", page=current, total=total, pct=int(current / total * 100))

    async def _ai_structure_resume(self, text: str) -> Optional[dict]:
        gateway = get_gateway()
        prompt = f"""Extract structured information from this resume text.

Return ONLY valid JSON (no markdown, no code fences) with this exact structure:
{{
  "name": "Full Name",
  "email": "email@example.com",
  "phone": "+1-555-555-5555",
  "skills": ["Skill 1", "Skill 2"],
  "experience": [
    {{"title": "Job Title", "organization": "Company", "date": "Start - End", "description": "Description", "highlights": ["achievement 1"]}}
  ],
  "projects": [
    {{"title": "Project Name", "organization": "", "date": "", "description": "Description", "highlights": ["detail 1"]}}
  ],
  "education": [
    {{"title": "Degree", "organization": "School", "date": "Year", "description": "", "highlights": []}}
  ],
  "certifications": ["Cert 1", "Cert 2"],
  "summary": "Brief professional summary"
}}

Resume text:
{text}"""

        response = await gateway.generate(
            prompt,
            system_prompt="You are a resume parser that extracts structured data. Return ONLY valid JSON.",
            temperature=0.1,
            max_tokens=4096,
        )

        cleaned = response.content.strip()
        if cleaned.startswith("```json"):
            cleaned = cleaned[7:]
        if cleaned.startswith("```"):
            cleaned = cleaned[3:]
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]

        try:
            parsed = json.loads(cleaned.strip())
            validated = {
                "name": parsed.get("name", ""),
                "email": parsed.get("email", ""),
                "phone": parsed.get("phone", ""),
                "skills": parsed.get("skills", []),
                "experience": parsed.get("experience", []),
                "projects": parsed.get("projects", []),
                "education": parsed.get("education", []),
                "certifications": parsed.get("certifications", []),
                "summary": parsed.get("summary", ""),
            }
            logger.info("ai_structure_success", has_name=bool(validated["name"]), skill_count=len(validated["skills"]))
            return validated
        except (json.JSONDecodeError, Exception) as e:
            logger.warning("ai_structure_parse_failed", error=str(e), response_snippet=cleaned[:200])
            return None

    def _fallback_parse(self, text: str) -> dict:
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
        phone_pattern = r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}"

        email = ""
        phone = ""
        name = lines[0] if lines else ""
        if len(name) > 60:
            name = ""

        for line in lines[:20]:
            if not email:
                found = re.findall(email_pattern, line)
                if found:
                    email = found[0]
            if not phone:
                found = re.findall(phone_pattern, line)
                if found:
                    phone = found[0]

        sections = {
            "skills": [], "experience": [], "education": [],
            "projects": [], "certifications": [], "summary": "",
        }
        current_section = "summary"
        section_keywords = {
            "skills": ["skills", "technologies", "tools", "technical", "competencies", "expertise"],
            "experience": ["experience", "work", "employment", "professional", "career", "history"],
            "education": ["education", "academic", "university", "college", "degree", "school", "bachelor", "master", "phd"],
            "projects": ["projects", "project"],
            "certifications": ["certifications", "certificates", "licenses"],
        }

        for line in lines:
            line_lower = line.lower()
            found_section = False
            for section, keywords in section_keywords.items():
                stripped = line_lower.strip(": ")
                if any(stripped == kw or stripped.startswith(kw) for kw in keywords):
                    current_section = section
                    found_section = True
                    break
            if found_section:
                continue
            if current_section == "skills":
                items = [s.strip() for s in re.split(r"[|,•·]", line) if s.strip()]
                sections["skills"].extend(items)
            elif current_section in ("experience", "education", "projects"):
                sections[current_section].append({"title": line, "organization": "", "date": "", "description": "", "highlights": []})
            elif current_section == "certifications":
                sections["certifications"].append(line)
            elif current_section == "summary":
                if sections["summary"]:
                    sections["summary"] += " " + line
                else:
                    sections["summary"] = line

        logger.info("fallback_parse_complete", skill_count=len(sections["skills"]))
        return {
            "name": name,
            "email": email,
            "phone": phone,
            "skills": sections["skills"],
            "experience": sections["experience"],
            "projects": sections["projects"],
            "education": sections["education"],
            "certifications": sections["certifications"],
            "summary": sections["summary"],
        }

    async def _generate_embedding(self, text: str) -> Optional[dict]:
        if not text or len(text.strip()) < 20:
            return None
        try:
            import hashlib
            text_hash = hashlib.sha256(text.encode()).hexdigest()
            import numpy as np
            fake_embedding = [float(hash(f"{text_hash}_{i}") % 1000) / 1000.0 for i in range(128)]
            return {"vector": fake_embedding, "dim": 128, "hash": text_hash}
        except Exception as e:
            logger.warning("embedding_failed", error=str(e))
            return None

    async def list_resumes(self, user_id: uuid.UUID) -> list:
        result = await self.db.execute(
            select(Resume)
            .where(Resume.user_id == user_id)
            .order_by(desc(Resume.created_at))
        )
        resumes = result.scalars().all()
        return [_resume_to_response(r) for r in resumes]

    async def get_resume(self, user_id: uuid.UUID, resume_id: str) -> ResumeResponse:
        result = await self.db.execute(
            select(Resume).where(
                Resume.id == resume_id,
                Resume.user_id == user_id,
            )
        )
        resume = result.scalar_one_or_none()
        if not resume:
            raise NotFoundException("Resume not found")
        return _resume_to_response(resume)

    async def update_resume(self, user_id: uuid.UUID, resume_id: str, parsed_data: dict) -> ResumeResponse:
        result = await self.db.execute(
            select(Resume).where(
                Resume.id == resume_id,
                Resume.user_id == user_id,
            )
        )
        resume = result.scalar_one_or_none()
        if not resume:
            raise NotFoundException("Resume not found")
        resume.parsed_data = parsed_data
        resume.parsing_confidence = 0.95
        await self.db.flush()
        await self.db.refresh(resume)
        return _resume_to_response(resume)

    async def delete_resume(self, user_id: uuid.UUID, resume_id: str):
        result = await self.db.execute(
            select(Resume).where(
                Resume.id == resume_id,
                Resume.user_id == user_id,
            )
        )
        resume = result.scalar_one_or_none()
        if not resume:
            raise NotFoundException("Resume not found")

        try:
            await delete_file(settings.S3_BUCKET, resume.file_key)
        except Exception as e:
            logger.warning("minio_delete_failed", file_key=resume.file_key, error=str(e))

        user_result = await self.db.execute(select(User).where(User.id == user_id))
        user = user_result.scalar_one()
        if user.active_resume_id == resume.id:
            fallback_result = await self.db.execute(
                select(Resume.id)
                .where(
                    Resume.user_id == user_id,
                    Resume.id != resume.id,
                    Resume.is_active == True,
                    Resume.parsing_status == ParsingStatus.COMPLETED,
                )
                .order_by(desc(Resume.created_at))
                .limit(1)
            )
            user.active_resume_id = fallback_result.scalar_one_or_none()

        await self.db.delete(resume)
        await self.db.flush()

    async def analyze_resume(self, user_id: uuid.UUID, resume_id: str) -> ResumeAnalysisResponse:
        result = await self.db.execute(
            select(Resume).where(
                Resume.id == resume_id,
                Resume.user_id == user_id,
            )
        )
        resume = result.scalar_one_or_none()
        if not resume:
            raise NotFoundException("Resume not found")

        parsed = resume.parsed_data or {}
        resume_text = json.dumps(parsed, indent=2)

        gateway = get_gateway()
        result_data = await gateway.analyze_resume(resume_text)

        skills_list = []
        for skill_name in parsed.get("skills", []):
            skills_list.append(SkillItem(name=skill_name, category="skill", proficiency="unknown", years=0))

        return ResumeAnalysisResponse(
            skills=skills_list,
            career_level=result_data.get("career_level"),
            industry=result_data.get("industry"),
            missing_skills=result_data.get("missing_skills", []),
            strengths=result_data.get("strengths", []),
            weaknesses=result_data.get("weaknesses", []),
            summary=result_data.get("summary"),
            experience_summary=result_data.get("experience_summary"),
            recommended_roles=result_data.get("recommended_roles", []),
            skill_graph=result_data.get("skill_graph"),
        )

    async def tailor_resume(
        self,
        user_id: uuid.UUID,
        resume_id: str,
        job_id: str,
        format_type: str = "ats_friendly",
    ) -> ResumeTailorResponse:
        result = await self.db.execute(
            select(Resume).where(
                Resume.id == resume_id,
                Resume.user_id == user_id,
            )
        )
        resume = result.scalar_one_or_none()
        if not resume:
            raise NotFoundException("Resume not found")

        job_result = await self.db.execute(select(Job).where(Job.id == job_id))
        job = job_result.scalar_one_or_none()

        parsed = resume.parsed_data or {}
        resume_text = json.dumps(parsed, indent=2)
        job_text = json.dumps({
            "title": job.title if job else "",
            "description": job.description if job else "",
            "skills": job.skills_required if job else [],
        }, indent=2)

        gateway = get_gateway()
        prompt = f"""Tailor this resume for the following job:

Resume data:
{resume_text}

Job details:
{job_text}

Format as {'ATS-friendly' if format_type == 'ats_friendly' else format_type} resume.

Return ONLY valid JSON with this structure (no markdown):
{{
  "sections": [
    {{"type": "header", "content": "candidate name and contact"}},
    {{"type": "summary", "content": "ATS-optimized professional summary"}},
    {{"type": "experience", "items": ["bullet point 1", "bullet point 2"]}},
    {{"type": "education", "items": ["degree, school, year"]}},
    {{"type": "skills", "items": ["skill1", "skill2"]}}
  ],
  "ats_score": 0-100
}}"""
        response = await gateway.generate(prompt, system_prompt="You are an expert ATS resume optimizer.")
        try:
            cleaned = response.content.strip()
            if cleaned.startswith("```json"):
                cleaned = cleaned[7:]
            if cleaned.startswith("```"):
                cleaned = cleaned[3:]
            if cleaned.endswith("```"):
                cleaned = cleaned[:-3]
            parsed_response = json.loads(cleaned.strip())
        except Exception:
            parsed_response = {"sections": [], "ats_score": 50}

        content = parsed_response.get("sections", [])
        ats_score = parsed_response.get("ats_score", 50.0)

        tailored = GeneratedResume(
            application_id=uuid.uuid4(),
            resume_id=uuid.UUID(resume_id),
            content={"sections": content},
            version=1,
            ai_generated=True,
            ats_score=float(ats_score),
        )
        self.db.add(tailored)
        await self.db.flush()

        return ResumeTailorResponse(
            id=str(tailored.id),
            ats_score=tailored.ats_score,
            version=tailored.version,
            content={"sections": content},
            created_at=tailored.created_at,
        )

    async def get_presigned_download_url(self, user_id: uuid.UUID, resume_id: str) -> str:
        result = await self.db.execute(
            select(Resume).where(
                Resume.id == resume_id,
                Resume.user_id == user_id,
            )
        )
        resume = result.scalar_one_or_none()
        if not resume:
            raise NotFoundException("Resume not found")
        try:
            url = await get_presigned_url(settings.S3_BUCKET, resume.file_key, expires_in=3600)
            return url
        except Exception as e:
            raise NotFoundException(f"Cannot access file: {str(e)}")


def _resume_to_response(r) -> ResumeResponse:
    return ResumeResponse(
        id=str(r.id),
        file_name=r.file_name,
        file_size=r.file_size,
        file_type=r.file_type.value if hasattr(r.file_type, "value") else str(r.file_type),
        raw_text=r.raw_text,
        parsed_data=r.parsed_data,
        parsing_confidence=r.parsing_confidence,
        parsing_status=r.parsing_status.value if hasattr(r.parsing_status, "value") else str(r.parsing_status),
        parsing_error=r.parsing_error,
        is_active=r.is_active,
        created_at=r.created_at,
        updated_at=r.updated_at,
    )
